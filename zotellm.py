"""
zotellm.py

One-command citation formatting: takes any document (.md or .docx) with informal
citations, resolves them via CrossRef/PubMed, and produces a Word document with
live Zotero field codes ready to refresh.

Usage:
    # From markdown
    python zotellm.py paper.md --provider cli --zotero-db ~/Zotero/zotero.sqlite

    # From Word document
    python zotellm.py paper.docx --provider cli --zotero-db ~/Zotero/zotero.sqlite

    # Output: paper_zotero.docx (open in Word, click Zotero > Refresh)

Requirements:
    pip install python-docx requests
"""

import argparse
import copy
import json
import os
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import time
from pathlib import Path

import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import shared functions from the other modules
# (we inline the key ones to keep this self-contained)

CROSSREF_API = "https://api.crossref.org/works"
CROSSREF_MAILTO = "crossref@example.com"
ZOTERO_API = "https://api.zotero.org"
SCHEMA_URL = (
    "https://github.com/citation-style-language/schema/raw/master/csl-citation.json"
)

PROVIDER_DEFAULTS = {
    "openai": "gpt-4o",
    "anthropic": "claude-sonnet-4-20250514",
}

# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------

CITATION_EXTRACTION_PROMPT = """\
You are a reference formatting assistant. Analyze the following document text \
and identify every citation or reference to a published work.

For each citation found, extract:
1. The text as it appears in the document (e.g., "Banwell et al., 2023" or "(Smith 2020)")
2. First author last name, year
3. **CRITICAL: Infer the likely title or topic of the cited work from the surrounding \
context.** For example, if the text says "The RAND/UCLA Appropriateness Method recommends \
a panel of 9 members (Fitch et al. 2001)", the title_hint should be "RAND UCLA \
appropriateness method manual". Be specific — this is the primary field used for search.
4. Any journal name mentioned or inferable from context
5. Any DOI or PMID if mentioned (e.g., "PMID: 12345678" or "doi: 10.1000/xyz")
6. A suggested citation key in the format: firstauthorlastnameYEAR (lowercase, no spaces)

Also identify any numbered reference list at the end and extract metadata from those entries.

Return a JSON object with this structure:
{
  "citations": [
    {
      "original_text": "Banwell et al., 2023",
      "context": "the sentence where this citation appears",
      "first_author": "Banwell",
      "year": "2023",
      "title_hint": "inferred title or topic keywords for search",
      "journal_hint": "any journal you can infer",
      "doi": "DOI if mentioned, otherwise empty string",
      "pmid": "PMID if mentioned, otherwise empty string",
      "suggested_key": "banwell2023"
    }
  ],
  "reference_list": [
    {
      "original_text": "1. Banwell B, Bennett JL, ...",
      "first_author": "Banwell",
      "year": "2023",
      "title": "full title if available",
      "journal": "The Lancet Neurology",
      "volume": "22",
      "pages": "268-282",
      "doi": "10.1016/...",
      "pmid": "PMID if available",
      "suggested_key": "banwell2023"
    }
  ]
}

Return ONLY the JSON object, no other text.

Document:
"""

DOCX_REWRITE_PROMPT = """\
You are a reference formatting assistant. Given a list of citation mappings, \
produce a JSON array of find-and-replace operations to convert informal citations \
to pandoc [@citekey] format.

Citation key mappings:
{mappings}

For each mapping, provide:
- "find": the exact text to find in the document (be precise — include parentheses if present)
- "replace": the [@citekey] replacement

Rules:
- For "(Author et al., Year)" → replace entire parenthetical with [@citekey]
- For "Author et al. (Year)" → replace with "Author et al. [@citekey]" only if the author name is part of the sentence flow, otherwise just [@citekey]
- For "(Author et al., Journal, Year)" → replace entire parenthetical with [@citekey]
- If the same citation appears multiple times, include one entry (it will be applied globally)

Return a JSON array like:
[
  {{"find": "(Banwell et al., 2023)", "replace": "[@banwell2023]"}},
  {{"find": "(Smith 2020)", "replace": "[@smith2020]"}}
]

Return ONLY the JSON array, no other text.
"""

MD_REWRITE_PROMPT = """\
You are a reference formatting assistant. Rewrite the following markdown document, \
replacing every inline citation with the pandoc citation syntax [@citekey].

Use these citation key mappings (original text -> citekey):
{mappings}

Rules:
- Replace "(Author et al., Year)" or "Author et al. (Year)" with [@citekey]
- For citations at the end of a sentence, place [@citekey] before the period
- For parenthetical citations, replace the entire parenthetical with [@citekey]
- If multiple citations are grouped, use [@key1; @key2]
- Remove any numbered reference list at the end (it will be auto-generated)
- Keep all other content exactly the same
- Preserve all YAML frontmatter, headings, figures, tables, etc.

Return ONLY the rewritten markdown, no explanation.

Document:
"""


# ---------------------------------------------------------------------------
# LLM backends
# ---------------------------------------------------------------------------

def llm_call(prompt, provider, model, api_base=None, api_key=None, max_tokens=8192,
             cli_command=None):
    if provider == "anthropic":
        return _call_anthropic(prompt, model, api_key, max_tokens)
    elif provider == "openai":
        return _call_openai(prompt, model, api_base, api_key, max_tokens)
    elif provider == "cli":
        return _call_cli(prompt, cli_command)
    else:
        print(f"Error: unknown provider '{provider}'.")
        sys.exit(1)


def _call_openai(prompt, model, api_base=None, api_key=None, max_tokens=8192):
    base = (api_base or os.environ.get("OPENAI_API_BASE", "https://api.openai.com/v1")).rstrip("/")
    key = api_key or os.environ.get("OPENAI_API_KEY", "")
    headers = {"Content-Type": "application/json"}
    if key:
        headers["Authorization"] = f"Bearer {key}"
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
        "temperature": 0,
    }
    resp = requests.post(f"{base}/chat/completions", headers=headers, json=payload, timeout=120)
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"].strip()


def _call_anthropic(prompt, model, api_key=None, max_tokens=8192):
    key = api_key or os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        print("Error: ANTHROPIC_API_KEY not set")
        sys.exit(1)
    headers = {
        "Content-Type": "application/json",
        "x-api-key": key,
        "anthropic-version": "2023-06-01",
    }
    payload = {"model": model, "max_tokens": max_tokens, "messages": [{"role": "user", "content": prompt}]}
    resp = requests.post("https://api.anthropic.com/v1/messages", headers=headers, json=payload, timeout=120)
    resp.raise_for_status()
    return resp.json()["content"][0]["text"].strip()


def _call_cli(prompt, cli_command=None):
    if cli_command:
        cmd = cli_command
    elif shutil.which("claude"):
        cmd = "claude --print"
    elif shutil.which("ollama"):
        cmd = "ollama run llama3"
    elif shutil.which("llm"):
        cmd = "llm"
    else:
        print("Error: no LLM CLI tool found. Install claude, ollama, or llm,")
        print("  or specify --cli-command 'your-command'")
        sys.exit(1)
    print(f"  Using CLI: {cmd}")
    env = os.environ.copy()
    env.pop("CLAUDECODE", None)
    result = subprocess.run(cmd, shell=True, input=prompt, capture_output=True,
                            text=True, timeout=300, env=env)
    if result.returncode != 0:
        print(f"CLI error (exit {result.returncode}): {result.stderr[:500]}")
        sys.exit(1)
    return result.stdout.strip()


# ---------------------------------------------------------------------------
# CrossRef / PubMed
# ---------------------------------------------------------------------------

def search_crossref(query, author=None, year=None, rows=5):
    params = {"query": query, "rows": rows, "mailto": CROSSREF_MAILTO}
    if author:
        params["query.author"] = author
    if year:
        params["query.bibliographic"] = f"{query} {year}"
    try:
        resp = requests.get(CROSSREF_API, params=params, timeout=15)
        resp.raise_for_status()
        return resp.json().get("message", {}).get("items", [])
    except Exception as e:
        print(f"  CrossRef search failed: {e}")
        return []


def crossref_to_csl(item):
    csl = {}
    csl["type"] = item.get("type", "article-journal").replace("journal-article", "article-journal")
    csl["title"] = item.get("title", [""])[0] if isinstance(item.get("title"), list) else item.get("title", "")
    csl["DOI"] = item.get("DOI", "")
    authors = []
    for a in item.get("author", []):
        au = {}
        if "family" in a:
            au["family"] = a["family"]
        if "given" in a:
            au["given"] = a["given"]
        if au:
            authors.append(au)
    if authors:
        csl["author"] = authors
    ct = item.get("container-title", [])
    if ct:
        csl["container-title"] = ct[0] if isinstance(ct, list) else ct
    for field in ("volume", "issue", "page"):
        if item.get(field):
            csl[field] = item[field]
    issued = item.get("issued", {})
    dp = issued.get("date-parts", [[]])
    if dp and dp[0]:
        csl["issued"] = {"date-parts": [dp[0]]}
    issn = item.get("ISSN", [])
    if issn:
        csl["ISSN"] = issn[0] if isinstance(issn, list) else issn
    return csl


def score_crossref_match(item, author=None, year=None, title_hint=None):
    score = 0
    if author and item.get("author"):
        fa = item["author"][0].get("family", "").lower()
        if fa == author.lower():
            score += 3
        elif author.lower() in fa:
            score += 1
    issued = item.get("issued", {}).get("date-parts", [[]])
    if issued and issued[0] and year:
        if str(issued[0][0]) == str(year):
            score += 3
    if title_hint:
        it = (item.get("title", [""])[0] if isinstance(item.get("title"), list)
              else item.get("title", "")).lower()
        hw = [w for w in title_hint.lower().split() if len(w) > 4]
        score += min(sum(1 for w in hw if w in it), 3)
    cr = item.get("score", 0)
    if cr > 100:
        score += 2
    elif cr > 50:
        score += 1
    return score


def search_pubmed(query, author=None, year=None, max_results=3):
    terms = []
    if author:
        terms.append(f"{author}[Author]")
    if year:
        terms.append(f"{year}[Date - Publication]")
    if query:
        terms.append(query)
    try:
        resp = requests.get("https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi",
                            params={"db": "pubmed", "term": " AND ".join(terms),
                                    "retmax": max_results, "retmode": "json"}, timeout=15)
        resp.raise_for_status()
        ids = resp.json().get("esearchresult", {}).get("idlist", [])
        if not ids:
            return []
        resp = requests.get("https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi",
                            params={"db": "pubmed", "id": ",".join(ids), "retmode": "json"}, timeout=15)
        resp.raise_for_status()
        result = resp.json().get("result", {})
        dois = []
        for pmid in ids:
            for aid in result.get(pmid, {}).get("articleids", []):
                if aid.get("idtype") == "doi":
                    dois.append(aid["value"])
                    break
        return dois
    except Exception:
        return []


def crossref_by_doi(doi):
    try:
        resp = requests.get(f"{CROSSREF_API}/{doi}", params={"mailto": CROSSREF_MAILTO}, timeout=15)
        resp.raise_for_status()
        return resp.json().get("message", {})
    except Exception:
        return None


def pmid_to_doi(pmid):
    """Look up a DOI from a PubMed ID."""
    try:
        resp = requests.get("https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi",
                            params={"db": "pubmed", "id": str(pmid), "retmode": "json"}, timeout=15)
        resp.raise_for_status()
        result = resp.json().get("result", {}).get(str(pmid), {})
        for aid in result.get("articleids", []):
            if aid.get("idtype") == "doi":
                return aid["value"]
    except Exception:
        pass
    return None


def find_best_match(citation):
    author = citation.get("first_author", "")
    year = citation.get("year", "")
    title = citation.get("title_hint") or citation.get("title", "")
    journal = citation.get("journal_hint") or citation.get("journal", "")
    queries = []
    if title and len(title) > 10:
        queries.append(title)
    if author and title and len(title) > 5:
        queries.append(f"{author} {title}")
    if author and year:
        queries.append(f"{author} {year}")
    if author and journal:
        queries.append(f"{author} {journal} {year}")
    best_item, best_score = None, -1
    for q in queries[:3]:
        for item in search_crossref(q, author=author, year=year):
            s = score_crossref_match(item, author, year, title)
            if s > best_score:
                best_score, best_item = s, item
        if best_score >= 6:
            break
        time.sleep(0.5)
    if best_score < 5 and (author or title):
        pq = title if title and len(title) > 10 else f"{author} {year}"
        for doi in search_pubmed(pq, author=author, year=year):
            item = crossref_by_doi(doi)
            if item:
                s = score_crossref_match(item, author, year, title)
                if s > best_score:
                    best_score, best_item = s, item
        time.sleep(0.5)
    return (best_item, best_score) if best_score >= 4 else (None, best_score)


# ---------------------------------------------------------------------------
# Zotero
# ---------------------------------------------------------------------------

def lookup_zotero_key_local(zotero_db, title=None, doi=None):
    if not zotero_db or not Path(zotero_db).exists():
        return None
    try:
        db = sqlite3.connect(str(zotero_db))
        if doi:
            row = db.execute("""
                SELECT i.key FROM items i
                JOIN itemData id ON i.itemID = id.itemID
                JOIN itemDataValues idv ON id.valueID = idv.valueID
                JOIN fields f ON id.fieldID = f.fieldID
                WHERE f.fieldName = 'DOI' AND LOWER(idv.value) = LOWER(?)
            """, (doi,)).fetchone()
            if row:
                db.close()
                return row[0]
        if title:
            row = db.execute("""
                SELECT i.key FROM items i
                JOIN itemData id ON i.itemID = id.itemID
                JOIN itemDataValues idv ON id.valueID = idv.valueID
                JOIN fields f ON id.fieldID = f.fieldID
                WHERE f.fieldName = 'title' AND LOWER(idv.value) LIKE LOWER(?)
            """, (f"%{title[:50]}%",)).fetchone()
            if row:
                db.close()
                return row[0]
        db.close()
    except Exception:
        pass
    return None


def get_zotero_item_id(zotero_db, zotero_key):
    import random
    if not zotero_db or not zotero_key:
        return random.randint(90000, 99999)
    try:
        db = sqlite3.connect(str(zotero_db))
        row = db.execute("SELECT itemID FROM items WHERE key = ?", (zotero_key,)).fetchone()
        db.close()
        return row[0] if row else random.randint(90000, 99999)
    except Exception:
        return random.randint(90000, 99999)


def get_zotero_user_id(zotero_db):
    if not zotero_db:
        return "0"
    try:
        db = sqlite3.connect(str(zotero_db))
        row = db.execute(
            "SELECT value FROM settings WHERE setting = 'account' AND key = 'userID'"
        ).fetchone()
        db.close()
        return str(row[0]) if row else "0"
    except Exception:
        return "0"


def add_to_zotero(api_key, library_id, csl_item, library_type="user"):
    type_map = {
        "article-journal": "journalArticle", "book": "book",
        "chapter": "bookSection", "paper-conference": "conferencePaper",
        "report": "report", "thesis": "thesis",
    }
    zot_item = {"itemType": type_map.get(csl_item.get("type", ""), "journalArticle")}
    field_map = {"title": "title", "container-title": "publicationTitle",
                 "volume": "volume", "issue": "issue", "page": "pages", "DOI": "DOI"}
    for csl_f, zot_f in field_map.items():
        if csl_item.get(csl_f):
            zot_item[zot_f] = csl_item[csl_f]
    if csl_item.get("issued", {}).get("date-parts"):
        zot_item["date"] = "-".join(str(p) for p in csl_item["issued"]["date-parts"][0])
    creators = [{"creatorType": "author", "firstName": a.get("given", ""),
                 "lastName": a.get("family", "")} for a in csl_item.get("author", [])]
    if creators:
        zot_item["creators"] = creators
    headers = {"Zotero-API-Key": api_key, "Zotero-API-Version": "3", "Content-Type": "application/json"}
    try:
        resp = requests.post(f"{ZOTERO_API}/{library_type}s/{library_id}/items",
                             headers=headers, json=[zot_item], timeout=15)
        resp.raise_for_status()
        result = resp.json()
        if result.get("successful"):
            return list(result["successful"].values())[0].get("key")
    except Exception as e:
        print(f"  Zotero API error: {e}")
    return None


# ---------------------------------------------------------------------------
# Document reading
# ---------------------------------------------------------------------------

def extract_text_from_docx(docx_path):
    """Extract all paragraph text from a .docx file."""
    doc = Document(docx_path)
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)
    return "\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Zotero field code insertion (from zotero_field_insert.py)
# ---------------------------------------------------------------------------

def random_id(length=8):
    import random, string
    return "".join(random.choices(string.ascii_lowercase + string.digits, k=length))


def build_citation_json(cite_keys, bib, keymap, zotero_db, user_id):
    citation_items = []
    for ck in cite_keys:
        zot_key = keymap.get(ck)
        item_id = get_zotero_item_id(zotero_db, zot_key)
        uri_key = zot_key or random_id().upper()
        uri = f"http://zotero.org/users/{user_id}/items/{uri_key}"
        csl = dict(bib.get(ck, {}))
        csl["id"] = f"{user_id}/{uri_key}"
        citation_items.append({"id": item_id, "uris": [uri], "itemData": csl})
    return {
        "citationID": random_id(),
        "properties": {"formattedCitation": "", "plainCitation": "", "noteIndex": 0},
        "citationItems": citation_items,
        "schema": SCHEMA_URL,
    }


def make_superscript_rpr(font="Calibri", size_pt=11):
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font)
    fonts.set(qn("w:hAnsi"), font)
    rpr.append(fonts)
    half_pt = str(size_pt * 2)
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), half_pt)
        rpr.append(el)
    va = OxmlElement("w:vertAlign")
    va.set(qn("w:val"), "superscript")
    rpr.append(va)
    return rpr


def insert_zotero_field(parent, position, citation_json, display_text, rpr_factory):
    runs = []
    for step in ("begin", "instr", "separate", "text", "end"):
        r = OxmlElement("w:r")
        r.append(copy.deepcopy(rpr_factory()))
        if step == "begin":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "begin")
            r.append(fc)
        elif step == "instr":
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " ADDIN ZOTERO_ITEM CSL_CITATION " + json.dumps(citation_json)
            r.append(instr)
        elif step == "separate":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "separate")
            r.append(fc)
        elif step == "text":
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = display_text
            r.append(t)
        elif step == "end":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "end")
            r.append(fc)
        runs.append(r)
    for run in reversed(runs):
        parent.insert(position, run)


def insert_zotero_bibl(paragraph):
    for child in list(paragraph._p):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "hyperlink"):
            paragraph._p.remove(child)
    for step, content in [
        ("begin", None),
        ("instr", ' ADDIN ZOTERO_BIBL {"uncited":[],"omitted":[],"custom":[]} CSL_BIBLIOGRAPHY'),
        ("separate", None),
        ("text", "[Bibliography will be generated by Zotero. Click Refresh in the Zotero tab.]"),
        ("end", None),
    ]:
        r = OxmlElement("w:r")
        if step == "begin":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "begin")
            r.append(fc)
        elif step == "instr":
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = content
            r.append(instr)
        elif step == "separate":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "separate")
            r.append(fc)
        elif step == "text":
            t = OxmlElement("w:t")
            t.text = content
            r.append(t)
        elif step == "end":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "end")
            r.append(fc)
        paragraph._p.append(r)


def process_paragraph_zotero(para, bib, keymap, zotero_db, user_id, rpr_factory):
    """Find [@key] markers in a paragraph and replace with Zotero fields."""
    full_text = para.text
    if "[@" not in full_text:
        return False
    markers = list(re.finditer(r"\[@([\w-]+)\]", full_text))
    if not markers:
        return False
    p_elem = para._p
    run_data = []
    for child in list(p_elem):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "r":
            texts = child.findall(qn("w:t"))
            run_text = "".join(t.text or "" for t in texts)
            rpr_elem = child.find(qn("w:rPr"))
            run_data.append({"text": run_text,
                             "rpr": copy.deepcopy(rpr_elem) if rpr_elem is not None else None})
        elif tag == "hyperlink":
            for sub_r in child.findall(qn("w:r")):
                texts = sub_r.findall(qn("w:t"))
                run_text = "".join(t.text or "" for t in texts)
                rpr_elem = sub_r.find(qn("w:rPr"))
                run_data.append({"text": run_text,
                                 "rpr": copy.deepcopy(rpr_elem) if rpr_elem is not None else None})
    concat = "".join(rd["text"] for rd in run_data)
    markers = list(re.finditer(r"\[@([\w-]+)\]", concat))
    if not markers:
        return False
    default_rpr = None
    for rd in run_data:
        if rd["rpr"] is not None and rd["text"].strip():
            default_rpr = rd["rpr"]
            break
    for child in list(p_elem):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "hyperlink"):
            p_elem.remove(child)
    pos = 0
    for marker in markers:
        before = concat[pos:marker.start()]
        if before:
            r = OxmlElement("w:r")
            if default_rpr is not None:
                r.append(copy.deepcopy(default_rpr))
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = before
            r.append(t)
            p_elem.append(r)
        cite_key = marker.group(1)
        if cite_key not in bib:
            print(f"  WARNING: '{cite_key}' not in bibliography")
        cit_json = build_citation_json([cite_key], bib, keymap, zotero_db, user_id)
        insert_zotero_field(p_elem, len(list(p_elem)), cit_json, f"[{cite_key}]", rpr_factory)
        pos = marker.end()
    after = concat[pos:]
    if after:
        r = OxmlElement("w:r")
        if default_rpr is not None:
            r.append(copy.deepcopy(default_rpr))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = after
        r.append(t)
        p_elem.append(r)
    return True


def replace_citations_in_docx(doc, replacements):
    """Replace informal citation text with [@citekey] markers in a docx Document."""
    count = 0
    for para in doc.paragraphs:
        p_elem = para._p
        # Collect all run text
        run_elements = []
        for child in list(p_elem):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "r":
                run_elements.append(child)

        if not run_elements:
            continue

        # Get full paragraph text
        full_text = para.text

        # Check if any replacement applies
        new_text = full_text
        changed = False
        for repl in replacements:
            find_text = repl.get("find", "")
            replace_text = repl.get("replace", "")
            if find_text and find_text in new_text:
                new_text = new_text.replace(find_text, replace_text)
                changed = True

        if not changed:
            continue

        # Get run properties from first text-bearing run
        default_rpr = None
        for child in run_elements:
            rpr_elem = child.find(qn("w:rPr"))
            texts = child.findall(qn("w:t"))
            if rpr_elem is not None and any(t.text and t.text.strip() for t in texts):
                default_rpr = copy.deepcopy(rpr_elem)
                break

        # Remove old runs
        for child in list(p_elem):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("r", "hyperlink"):
                p_elem.remove(child)

        # Add single new run with replaced text
        r = OxmlElement("w:r")
        if default_rpr is not None:
            r.append(default_rpr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = new_text
        r.append(t)
        p_elem.append(r)
        count += 1

    return count


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def parse_json_response(raw):
    """Parse JSON from LLM response, handling markdown code blocks."""
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```\w*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)
    return json.loads(raw)


def main():
    parser = argparse.ArgumentParser(
        description="zotellm: one-command citation formatting with Zotero field codes"
    )
    parser.add_argument("input", help="Input file (.md or .docx) with informal citations")
    parser.add_argument("--output", "-o", help="Output .docx path (default: input_zotero.docx)")
    parser.add_argument("--provider", "-p", default="openai",
                        choices=["openai", "anthropic", "cli"],
                        help="LLM provider (default: openai). 'cli' uses claude/ollama/llm.")
    parser.add_argument("--model", "-m", help="Model name (default depends on provider)")
    parser.add_argument("--api-base", help="API base URL for custom endpoints")
    parser.add_argument("--api-key", help="API key (overrides env var)")
    parser.add_argument("--cli-command", help="Custom CLI command for --provider cli")
    parser.add_argument("--zotero-db", help="Path to local zotero.sqlite")
    parser.add_argument("--zotero-api-key", help="Zotero Web API key (for adding items)")
    parser.add_argument("--zotero-library-id", help="Zotero user library ID")
    parser.add_argument("--reference-doc", help="Pandoc reference .docx template (for .md input)")
    parser.add_argument("--font", default="Calibri", help="Font for citation text (default: Calibri)")
    parser.add_argument("--size", type=int, default=11, help="Font size in pt (default: 11)")
    parser.add_argument("--bib-heading", default="References",
                        help="Heading where bibliography should be inserted (default: References)")
    parser.add_argument("--no-crossref", action="store_true", help="Skip CrossRef lookups")
    parser.add_argument("--dry-run", action="store_true", help="Preview without writing files")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: {input_path} not found")
        sys.exit(1)

    is_docx = input_path.suffix.lower() == ".docx"
    is_md = input_path.suffix.lower() in (".md", ".markdown", ".txt")

    if not is_docx and not is_md:
        print(f"Error: unsupported file type '{input_path.suffix}'. Use .md or .docx.")
        sys.exit(1)

    output_path = args.output or str(input_path.with_suffix("")) + "_zotero.docx"
    model = args.model or PROVIDER_DEFAULTS.get(args.provider, "gpt-4o")
    llm_kwargs = dict(provider=args.provider, model=model, api_base=args.api_base,
                      api_key=args.api_key, cli_command=args.cli_command)

    print(f"Input: {input_path} ({'Word' if is_docx else 'Markdown'})")
    print(f"Provider: {args.provider}" + (f" | Model: {model}" if args.provider != "cli" else ""))

    # --- Step 1: Extract text ---
    if is_docx:
        text = extract_text_from_docx(input_path)
    else:
        text = input_path.read_text()

    # --- Step 2: Identify citations with LLM ---
    print("\nStep 1: Identifying citations...")
    raw = llm_call(CITATION_EXTRACTION_PROMPT + text, max_tokens=4096, **llm_kwargs)
    extracted = parse_json_response(raw)

    citations = extracted.get("citations", [])
    ref_list = extracted.get("reference_list", [])
    print(f"  Found {len(citations)} inline citations, {len(ref_list)} reference list entries")

    # Merge
    all_refs = {}
    for ref in ref_list:
        key = ref.get("suggested_key", "")
        if key:
            all_refs[key] = ref
    for cit in citations:
        key = cit.get("suggested_key", "")
        if key and key not in all_refs:
            all_refs[key] = cit
        elif key and key in all_refs:
            if not all_refs[key].get("title_hint") and cit.get("title_hint"):
                all_refs[key]["title_hint"] = cit["title_hint"]

    if not all_refs:
        print("  No citations found. Nothing to do.")
        return

    print(f"  {len(all_refs)} unique references to resolve")

    # --- Step 3: Resolve references ---
    print("\nStep 2: Resolving references...")
    bib_items = []
    keymap = {}
    mappings = []

    for key, ref in all_refs.items():
        author = ref.get("first_author", "")
        year = ref.get("year", "")
        print(f"  [{key}] {author} {year}...", end=" ")

        zotero_key = None
        if args.zotero_db:
            title = ref.get("title") or ref.get("title_hint", "")
            doi = ref.get("doi", "")
            zotero_key = lookup_zotero_key_local(args.zotero_db, title=title, doi=doi)
            if zotero_key:
                print(f"found in Zotero [{zotero_key}]")

        # Direct lookup by PMID or DOI if available
        csl = None
        ref_doi = ref.get("doi", "")
        ref_pmid = ref.get("pmid", "")

        if ref_doi and not args.no_crossref:
            cr_item = crossref_by_doi(ref_doi)
            if cr_item:
                csl = crossref_to_csl(cr_item)
                csl["id"] = key
                print(f"DOI match - {csl.get('title', '')[:60]}")
        elif ref_pmid and not args.no_crossref:
            doi_from_pmid = pmid_to_doi(ref_pmid)
            if doi_from_pmid:
                cr_item = crossref_by_doi(doi_from_pmid)
                if cr_item:
                    csl = crossref_to_csl(cr_item)
                    csl["id"] = key
                    print(f"PMID→DOI match - {csl.get('title', '')[:60]}")

        # Fall back to search if no direct match
        if csl is None and not args.no_crossref:
            cr_item, score = find_best_match(ref)
            if cr_item:
                csl = crossref_to_csl(cr_item)
                csl["id"] = key
                print(f"CrossRef match (score={score})" +
                      (f" - {csl.get('title', '')[:60]}" if not zotero_key else ""))
            else:
                print("no match")
        elif csl is None:
            print("skipping lookup")

        if csl is None:
            csl = {"id": key, "type": "article-journal",
                   "title": ref.get("title") or ref.get("title_hint", f"[{key}]")}
            if author:
                csl["author"] = [{"family": author}]
            if year:
                csl["issued"] = {"date-parts": [[int(year)]]}
            if ref.get("journal") or ref.get("journal_hint"):
                csl["container-title"] = ref.get("journal") or ref.get("journal_hint")

        bib_items.append(csl)
        keymap[key] = zotero_key

        if not zotero_key and args.zotero_api_key and args.zotero_library_id:
            print(f"    Adding to Zotero...", end=" ")
            new_key = add_to_zotero(args.zotero_api_key, args.zotero_library_id, csl)
            if new_key:
                keymap[key] = new_key
                print(f"added [{new_key}]")
            else:
                print("failed")

        orig = ref.get("original_text", "")
        if orig:
            mappings.append(f'"{orig}" -> [@{key}]')

    bib_dict = {item["id"]: item for item in bib_items}

    if args.dry_run:
        print(f"\n--- DRY RUN ---")
        print(f"References: {[i['id'] for i in bib_items]}")
        print(f"Keymap: {json.dumps(keymap, indent=2)}")
        return

    # --- Step 4: Build the output .docx ---
    user_id = get_zotero_user_id(args.zotero_db) if args.zotero_db else "0"
    rpr_factory = lambda: make_superscript_rpr(args.font, args.size)

    if is_docx:
        # For .docx input: replace inline citations, then insert Zotero fields
        print(f"\nStep 3: Replacing citations in Word document...")
        mapping_text = "\n".join(mappings)
        rewrite_prompt = DOCX_REWRITE_PROMPT.replace("{mappings}", mapping_text)
        raw = llm_call(rewrite_prompt, max_tokens=4096, **llm_kwargs)
        replacements = parse_json_response(raw)
        print(f"  {len(replacements)} replacements to apply")

        doc = Document(str(input_path))
        n_replaced = replace_citations_in_docx(doc, replacements)
        print(f"  Replaced citations in {n_replaced} paragraphs")

        print(f"\nStep 4: Inserting Zotero field codes...")
        n_fields = 0
        for para in doc.paragraphs:
            if process_paragraph_zotero(para, bib_dict, keymap, args.zotero_db, user_id, rpr_factory):
                n_fields += 1
        print(f"  Inserted Zotero fields in {n_fields} paragraphs")

        # Insert bibliography
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if re.match(rf"^(\d+\.\s*)?{re.escape(args.bib_heading)}$", text):
                if i + 1 < len(doc.paragraphs):
                    insert_zotero_bibl(doc.paragraphs[i + 1])
                else:
                    insert_zotero_bibl(doc.add_paragraph())
                print("  Inserted ZOTERO_BIBL field")
                break

        doc.save(output_path)

    else:
        # For .md input: rewrite markdown, convert with pandoc, insert Zotero fields
        print(f"\nStep 3: Rewriting markdown with [@citekey] markers...")
        mapping_text = "\n".join(mappings)
        rewrite_prompt = MD_REWRITE_PROMPT.replace("{mappings}", mapping_text)
        rewritten = llm_call(rewrite_prompt + text, max_tokens=8192, **llm_kwargs)
        if rewritten.startswith("```"):
            rewritten = re.sub(r"^```\w*\n?", "", rewritten)
            rewritten = re.sub(r"\n?```$", "", rewritten)

        # Write temp markdown
        tmp_md = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False)
        tmp_md.write(rewritten)
        tmp_md.close()

        # Convert with pandoc
        print(f"\nStep 4: Converting to .docx with pandoc...")
        pandoc_cmd = ["pandoc", tmp_md.name, "-o", output_path]
        if args.reference_doc:
            pandoc_cmd.extend(["--reference-doc", args.reference_doc])
        result = subprocess.run(pandoc_cmd, capture_output=True, text=True)
        os.unlink(tmp_md.name)
        if result.returncode != 0:
            print(f"Pandoc error: {result.stderr}")
            sys.exit(1)

        # Insert Zotero fields
        print(f"\nStep 5: Inserting Zotero field codes...")
        doc = Document(output_path)
        n_fields = 0
        for para in doc.paragraphs:
            if process_paragraph_zotero(para, bib_dict, keymap, args.zotero_db, user_id, rpr_factory):
                n_fields += 1
        print(f"  Inserted Zotero fields in {n_fields} paragraphs")

        for i, para in enumerate(doc.paragraphs):
            text_stripped = para.text.strip()
            if re.match(rf"^(\d+\.\s*)?{re.escape(args.bib_heading)}$", text_stripped):
                if i + 1 < len(doc.paragraphs):
                    insert_zotero_bibl(doc.paragraphs[i + 1])
                else:
                    insert_zotero_bibl(doc.add_paragraph())
                print("  Inserted ZOTERO_BIBL field")
                break

        doc.save(output_path)

    # Save bibliography and keymap alongside output
    bib_path = str(Path(output_path).with_suffix("")) + "_references.json"
    keymap_path = str(Path(output_path).with_suffix("")) + "_keymap.json"
    with open(bib_path, "w") as f:
        json.dump(bib_items, f, indent=2)
    with open(keymap_path, "w") as f:
        json.dump(keymap, f, indent=2)

    print(f"\nDone!")
    print(f"  Output: {output_path}")
    print(f"  Bibliography: {bib_path} ({len(bib_items)} items)")
    print(f"  Keymap: {keymap_path}")
    print(f"  Open in Word and click Zotero > Refresh")


if __name__ == "__main__":
    main()
